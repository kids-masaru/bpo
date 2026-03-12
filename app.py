import os
import json
import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
import hashlib
from google import genai
from google.genai import types
import openpyxl
from docx import Document as DocxDocument
import copy
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go


# --- Configuration & Constants ---
SERVICE_ACCOUNT_FILE = "service_account.json"
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets"
]

# JSON Schema for Full Data Extraction
GEMINI_SCHEMA = {
    "type": "object",
    "properties": {
        "B_都道府県": {"type": "string"},
        "C_自治体名": {"type": "string"},
        "D_案件名": {"type": "string"},
        "E_施設区分": {"type": "string", "enum": ["子育て支援拠点", "児童館", "学童", "放課後子供教室", "複合施設", "その他"]},
        "F_契約形態": {"type": "string", "enum": ["業務委託", "指定管理者", "補助金交付"]},
        "H_公示日": {"type": "string"},
        "I_参加表明期限": {"type": "string"},
        "J_提案書提出期限": {"type": "string"},
        "K_プレゼン実施日": {"type": "string"},
        "L_事業開始日": {"type": "string"},
        "M_契約期間年数": {"type": "integer"},
        "N_契約方式": {"type": "string"},
        "O_年間基本額上限": {"type": "integer"},
        "P_初期開設補助の有無": {"type": "string"},
        "Q_利用料徴収主体": {"type": "string", "enum": ["事業者が徴収し自社の収入とする", "事業者が徴収し市へ納入する", "市が直接徴収する", "完全無料"]},
        "R_光熱水費の負担": {"type": "string", "enum": ["市", "事業者", "上限付き事業者負担"]},
        "S_消耗品おやつ代負担": {"type": "string", "enum": ["市", "事業者", "実費徴収"]},
        "T_物価変動リスク": {"type": "string", "enum": ["市が補填", "事業者が吸収"]},
        "U_エリア指定": {"type": "string"},
        "V_必須配置人数": {"type": "integer"},
        "W_必須資格要件": {"type": "string"},
        "X_施設長要件": {"type": "string"},
        "Y_開所時間": {"type": "string"},
        "Z_独自コンテンツ要件": {"type": "string"},
        "AA_評価総点": {"type": "integer"},
        "AB_価格点割合パーセント": {"type": "number"},
        "AC_企画事業内容配点": {"type": "integer"},
        "AD_職員体制実績配点": {"type": "integer"},
        "AE_安全管理危機管理配点": {"type": "integer"},
        "AF_業務提供スタイル": {"type": "string", "enum": ["施設型", "訪問型", "一般型", "余裕活用型"]},
        "AG_送迎運行業務の有無": {"type": "boolean"},
        "AH_特殊な必須業務": {"type": "string"},
        "AI_プレゼン時間分": {"type": "integer"},
        "AJ_競合参入リスク": {"type": "string", "enum": ["高", "中", "低"]},
        "AK_人員確保難易度": {"type": "integer"},
        "AL_特記事項懸念点まとめ": {"type": "string"},
        "AM_必須資格フラグ": {"type": "boolean"},
        "AN_公募年度": {"type": "string"},
        "AO_提案要求リスト": {"type": "string"},
        "AP_物理的システム的準備リスト": {"type": "string"}
    }
}

# Schema for Quick Scan
SCAN_SCHEMA = {
    "type": "object",
    "properties": {
        "summary": {"type": "string"},
        "municipality": {"type": "string"},
        "project_name": {"type": "string"},
        "fiscal_year": {"type": "string"}
    }
}

# Schema for Grouping
GROUP_SCHEMA = {
    "type": "object",
    "properties": {
        "groups": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "group_name": {"type": "string"},
                    "municipality": {"type": "string"},
                    "project_name": {"type": "string"},
                    "fiscal_year": {"type": "string"},
                    "file_indices": {"type": "array", "items": {"type": "integer"}}
                }
            }
        }
    }
}

# Schema for AI Deduplication Matching
MATCH_SCHEMA = {
    "type": "object",
    "properties": {
        "matches": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "group_index": {"type": "integer"},
                    "matched_existing_row": {"type": "integer"}
                }
            }
        }
    }
}

# Column Layout (A ~ AN)
COLUMN_LAYOUT = [
    "A_No", "B_都道府県", "C_自治体名", "D_案件名", "E_施設区分", "F_契約形態", "G_ステータス",
    "H_公示日", "I_参加表明期限", "J_提案書提出期限", "K_プレゼン実施日", "L_事業開始日",
    "M_契約期間年数", "N_契約方式", "O_年間基本額上限", "P_初期開設補助の有無", "Q_利用料徴収主体",
    "R_光熱水費の負担", "S_消耗品おやつ代負担", "T_物価変動リスク", "U_エリア指定",
    "V_必須配置人数", "W_必須資格要件", "X_施設長要件", "Y_開所時間", "Z_独自コンテンツ要件",
    "AA_評価総点", "AB_価格点割合パーセント", "AC_企画事業内容配点", "AD_職員体制実績配点",
    "AE_安全管理危機管理配点", "AF_業務提供スタイル", "AG_送迎運行業務の有無", "AH_特殊な必須業務",
    "AI_プレゼン時間分", "AJ_競合参入リスク", "AK_人員確保難易度", "AL_特記事項懸念点まとめ",
    "AM_必須資格フラグ", "AN_公募年度", "AO_提案要求リスト", "AP_物理的システム的準備リスト"
]

# Display-friendly column names (without prefix)
DISPLAY_COLS = {col: col.split("_", 1)[1] if "_" in col else col for col in COLUMN_LAYOUT}

COL_IDX_MUNICIPALITY = COLUMN_LAYOUT.index("C_自治体名")
COL_IDX_PROJECT = COLUMN_LAYOUT.index("D_案件名")
COL_IDX_YEAR = COLUMN_LAYOUT.index("AN_公募年度")

# --- Service Initializers ---
def get_service_account_info():
    if "gcp_service_account" in st.secrets:
        return dict(st.secrets["gcp_service_account"])
    elif os.path.exists(SERVICE_ACCOUNT_FILE):
        with open(SERVICE_ACCOUNT_FILE) as f:
            return json.load(f)
    else:
        st.error("Google Cloud の認証情報が見つかりません。")
        st.stop()

def get_sheets_client():
    info = get_service_account_info()
    creds = Credentials.from_service_account_info(info, scopes=SCOPES)
    return gspread.authorize(creds)

def get_sheet(spreadsheet_id_or_name):
    client = get_sheets_client()
    if len(spreadsheet_id_or_name) > 30:
        return client.open_by_key(spreadsheet_id_or_name).worksheet("案件データ")
    else:
        return client.open(spreadsheet_id_or_name).worksheet("案件データ")

# --- File Conversion ---
GEMINI_SUPPORTED_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".txt", ".csv", ".md"}

def convert_file_for_gemini(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in GEMINI_SUPPORTED_EXTENSIONS and ext not in (".txt", ".md", ".csv"):
        # Image or PDF files are uploaded directly (Gemini handles their tokenization internally usually better)
        return file_path, []
    
    txt_path = file_path + ".converted.txt"
    text_content = ""
    
    try:
        if ext in (".xlsx", ".xls"):
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_content += f"\n=== シート: {sheet_name} ===\n"
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        text_content += "\t".join(cells) + "\n"
            wb.close()
        elif ext in (".docx",):
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text_content += "\t".join(cells) + "\n"
                text_content += "\n"
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
            except Exception:
                text_content = f"（このファイル形式 [{ext}] は読み取れませんでした）"
    except Exception as e:
        text_content = f"（ファイル変換エラー: {str(e)}）"
    
    # --- 安全装置: 巨大すぎるファイルは文字数をカットしてトークン上限エラーを防ぐ ---
    # 日本語の100万トークンは約100万文字〜200万文字。安全マージンを取って50万文字で切る。
    MAX_CHARS = 500000
    if len(text_content) > MAX_CHARS:
        text_content = text_content[:MAX_CHARS] + f"\n\n（※ファイルが巨大すぎるため、最初の{MAX_CHARS}文字のみを抽出し以降をカットしました）"
    
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text_content)
    return txt_path, [txt_path]

def upload_file_to_gemini(client, file_path):
    converted_path, temp_files = convert_file_for_gemini(file_path)
    file_upload = client.files.upload(file=converted_path)
    for tf in temp_files:
        if os.path.exists(tf):
            os.remove(tf)
    return file_upload

# --- AI Functions ---

def step1_scan_files(api_key, file_paths, status_callback=None):
    client = genai.Client(api_key=api_key)
    summaries = []
    for i, path in enumerate(file_paths):
        if status_callback:
            status_callback(f"🔍 Step 1: ファイル {i+1}/{len(file_paths)} をスキャン中...")
        file_upload = upload_file_to_gemini(client, path)
        prompt = """この資料の内容をスキャンして以下を特定してください。
・summary: 内容の一行要約
・municipality: 自治体の名前（市区町村名）
・project_name: 案件名・事業名
・fiscal_year: 公募の年度（例: 令和8年度、2025年度）。日付から推測してください。"""
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=[prompt, file_upload],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=SCAN_SCHEMA,
                temperature=0.1,
            )
        )
        scan_result = json.loads(response.text)
        scan_result["file_index"] = i
        scan_result["file_path"] = path
        summaries.append(scan_result)
    return summaries

def step2_group_files(api_key, summaries):
    client = genai.Client(api_key=api_key)
    summary_text = ""
    for s in summaries:
        summary_text += (
            f"ファイル{s['file_index']}: "
            f"自治体={s.get('municipality', '不明')}, "
            f"案件={s.get('project_name', '不明')}, "
            f"年度={s.get('fiscal_year', '不明')}, "
            f"概要={s.get('summary', '')}\n"
        )
    prompt = f"""以下は複数の資料のスキャン結果です。これらを「同じ自治体の同じ案件の同じ年度に関する資料」ごとにグループ分けしてください。

【重要ルール】
・同じ自治体 + 同じ案件 + 同じ年度 → 同じグループ
・同じ自治体 + 同じ案件 + 違う年度 → 別グループ
・仕様書・公告・評価基準表・契約書・申請書様式など、同じ公募パッケージに含まれる書類は全て同じグループ
・迷った場合は、同じグループにまとめる方向で判断してください（分け過ぎ厳禁）

{summary_text}

各グループに分かりやすい名前を付けてください。"""
    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=[prompt],
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=GROUP_SCHEMA,
            temperature=0.1,
        )
    )
    return json.loads(response.text)

def step3_check_existing(api_key, sheet, groups_list):
    try:
        all_rows = sheet.get_all_values()
    except Exception:
        all_rows = []
    
    # Get the last 20 rows (excluding header)
    data_rows = all_rows[1:] if len(all_rows) > 1 else []
    recent_rows = data_rows[-20:] if len(data_rows) > 0 else []
    
    # Calculate actual row numbers for the recent rows (1-indexed, +1 for header)
    start_row_idx = len(data_rows) - len(recent_rows) + 2 if data_rows else 2
    
    recent_list_text = ""
    for i, row in enumerate(recent_rows):
        actual_row_num = start_row_idx + i
        e_muni = row[COL_IDX_MUNICIPALITY].strip() if len(row) > COL_IDX_MUNICIPALITY else ""
        e_proj = row[COL_IDX_PROJECT].strip() if len(row) > COL_IDX_PROJECT else ""
        e_year = row[COL_IDX_YEAR].strip() if len(row) > COL_IDX_YEAR else ""
        recent_list_text += f"[行番号: {actual_row_num}] 自治体: {e_muni}, 案件名: {e_proj}, 年度: {e_year}\n"

    new_groups_text = ""
    for gi, group in enumerate(groups_list):
        new_groups_text += f"[新規グループ {gi}] 自治体: {group.get('municipality', '')}, 案件名: {group.get('project_name', '')}, 年度: {group.get('fiscal_year', '')}\n"

    annotated = []
    
    if not recent_list_text.strip() or not new_groups_text.strip():
        # If no existing data or no new groups, everything is new
        for group in groups_list:
            annotated.append({**group, "mode": "new", "existing_row": None})
        return annotated

    client = genai.Client(api_key=api_key)
    prompt = f"""以下は、スプレッドシートに登録済みの直近20件の案件データと、今回新しく読み取った案件のリストです。
新しく読み取った案件が、既存のデータと「同一案件」であるかを確認してください。

【判定ルール】
1. 和暦と西暦の違い（例：令和7年度 と 2025年度）は同一年度とみなす。
2. 「業務委託」「プロポーザル」「運営業務」といった末尾の単語の有無などの表記揺れは、文脈から判断して実質的に同じ案件であれば同一とみなす。
3. 一致する既存案件が見つかった場合は、その「行番号」を返してください。
4. 一致しない場合は、その新規グループについてはリストに含めない、または null として処理してください。
5. 【重要】「今回新しく読み取った案件」のすべて（新規グループ0〜全て）について、必ず1つずつ既存データと照らし合わせて確認し、漏れなく(面倒くさがらずに)全ての一致結果を配列に含めてください。

▼ 既存のデータ（直近20件）
{recent_list_text}

▼ 今回新しく読み取った案件
{new_groups_text}
"""
    try:
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=[prompt],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=MATCH_SCHEMA,
                temperature=0.1,
            )
        )
        match_result = json.loads(response.text)
        print(f"DEBUG - AI Match Result: {match_result}")
        
        # Build mapping from group_index to matched row
        matches_map = {}
        for match in match_result.get("matches", []):
            g_idx = match.get("group_index")
            e_row = match.get("matched_existing_row")
            if g_idx is not None and e_row is not None:
                matches_map[int(g_idx)] = int(e_row)
                
        for gi, group in enumerate(groups_list):
            matched_row = matches_map.get(gi)
            if matched_row:
                annotated.append({**group, "mode": "update", "existing_row": matched_row})
            else:
                annotated.append({**group, "mode": "new", "existing_row": None})
                
    except Exception as e:
        # Fallback to pure new if AI matching fails completely
        print(f"AI Deduplication failed: {e}")
        for group in groups_list:
            annotated.append({**group, "mode": "new", "existing_row": None})

    return annotated

def step4_extract_data(api_key, file_paths, existing_data=None):
    client = genai.Client(api_key=api_key)
    uploaded_files = []
    for path in file_paths:
        file_upload = upload_file_to_gemini(client, path)
        uploaded_files.append(file_upload)
    
    prompt = """提供された資料は、全て「1つの自治体の公募案件」に関連するものです。
これらを複合的に読み解き、統合して情報をJSON形式で抽出してください。
該当情報がない場合は null または空文字を設定してください。

【表記統一ルール（非常に重要）】
1. 日付の西暦化: 公示日、提出期限、事業開始日などのすべての日付は、資料に「令和7年4月1日」や「R7.4.1」と書かれていても、可能な限り「2025年4月1日」という形に西暦で統一して抽出してください（年は西暦、月日はそのまま）。
2. 数字の半角化: 人数、金額、期間などの数字は、すべて「半角アラビア数字（123...）」に統一してください。単位がある場合はその後ろにつけてください（例：３名 → 3名）。

【AI評価項目の判断基準】
・AJ_競合参入リスク: 「高」（参入障壁低く激戦） / 「中」 / 「低」（特殊要件多く他社入りにくい）
・AK_人員確保難易度: 1（容易） 〜 5（非常に困難）で評価。
・AL_特記事項懸念点まとめ: リスクや旨味、変わった要件などを100文字以内で。
・AM_必須資格フラグ: 有資格者が必須なら true そうでなければ false。
・AN_公募年度: 公募の年度（例: 令和8年度）。
・AO_提案要求リスト: 仕様書・評価基準内で「提案書に記載・提示しなさい」と求められている企画やテーマなどを全て漏らさず改行と中ポツ(・)の箇条書きで抽出。（例:\n・新規利用者の確保に向けた広報計画\n・保護者クレーム対応フロー）
・AP_物理的システム的準備リスト: 提案のアイデアではなく、実際に「導入・用意・採用」しなければならない物理的枠組み・金銭的・制度的な要件（ICT等）を全て漏らさず改行と中ポツ(・)の箇条書きで抽出。（例:\n・保護者連絡用ICTシステムの導入\n・常勤保育士3名の配置）
"""
    if existing_data:
        prompt += f"""
【重要】既存データがあります。新しい資料で補完・修正できる情報は更新し、既存データが正しい部分は維持してください。
既存データ: {json.dumps(existing_data, ensure_ascii=False)}
"""
    contents = [prompt] + uploaded_files
    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=contents,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=GEMINI_SCHEMA,
            temperature=0.2,
        )
    )
    return json.loads(response.text)

def step5_write_to_sheet(sheet, row_data, mode, existing_row=None):
    if mode == "update" and existing_row:
        col_count = len(COLUMN_LAYOUT)
        if col_count <= 26:
            end_col = chr(64 + col_count)
        else:
            end_col = chr(64 + (col_count - 1) // 26) + chr(64 + (col_count - 1) % 26 + 1)
        cell_range = f"A{existing_row}:{end_col}{existing_row}"
        sheet.update(range_name=cell_range, values=[row_data], value_input_option="USER_ENTERED")
    else:
        sheet.append_row(row_data, value_input_option="USER_ENTERED")
    return True

def build_row(extracted_data, row_number, mode="new", sheet=None, existing_row=None):
    row = []
    for col in COLUMN_LAYOUT:
        if col == "A_No":
            if mode == "update" and sheet and existing_row:
                try:
                    existing_no = sheet.cell(existing_row, 1).value
                    row.append(existing_no if existing_no else row_number)
                except Exception:
                    row.append(row_number)
            else:
                row.append(row_number)
        elif col == "G_ステータス":
            row.append("FALSE")
        else:
            val = extracted_data.get(col)
            row.append(val if val is not None else "")
    return row

# --- Dashboard Functions ---
@st.cache_data(ttl=600)
def load_historical_data(spreadsheet_id):
    try:
        sheet = get_sheet(spreadsheet_id)
        data = sheet.get_all_values()
        if not data:
            return pd.DataFrame()
        
        headers = data[0]
        df = pd.DataFrame(data[1:], columns=headers)
        
        # Fuzzy map natural language headers back to COLUMN_LAYOUT keys
        rename_dict = {}
        # Hardcoded fallbacks for headers that changed significantly in the sheet
        hardcoded_map = {
            "年間基本額（委託料/補助金額上限）": "O_年間基本額上限",
            "必須配置人数（常時）": "V_必須配置人数",
            "評価配点（総点）": "AA_評価総点",
            "年度": "AN_公募年度",
            "保護者からの利用料徴収主体": "Q_利用料徴収主体",
            "消耗品・おやつ代の負担": "S_消耗品おやつ代負担",
            "エリア指定・学区": "U_エリア指定",
            "施設長（責任者）要件": "X_施設長要件",
            "開所時間・延長保育": "Y_開所時間",
            "【配点】価格点（%）": "AB_価格点割合パーセント",
            "【配点】企画・事業内容": "AC_企画事業内容配点",
            "【配点】職員体制・実績": "AD_職員体制実績配点",
            "【配点】安全管理・危機管理": "AE_安全管理危機管理配点",
            "特殊な必須業務（キーワード抽出）": "AH_特殊な必須業務"
        }
        
        for h in df.columns:
            h_clean = str(h).strip()
            if h_clean == "No.":
                rename_dict[h] = "A_No"
                continue
            
            if h_clean in hardcoded_map:
                rename_dict[h] = hardcoded_map[h_clean]
                continue
                
            # Perform fuzzy match to handle brackets and spaces
            h_fuzzy = h_clean.replace("（", "").replace("）", "").replace("(", "").replace(")", "").replace("％", "").replace("%", "").replace("・", "")
            for col_key in COLUMN_LAYOUT:
                expected_desc = col_key.split("_", 1)[1] if "_" in col_key else col_key
                expected_fuzzy = expected_desc.replace("（", "").replace("）", "").replace("(", "").replace(")", "").replace("％", "").replace("%", "").replace("・", "")
                
                if expected_fuzzy == h_fuzzy:
                    rename_dict[h] = col_key
                    break
        df.rename(columns=rename_dict, inplace=True)
        
        # Ensure numeric columns are converted
        numeric_cols = ["M_契約期間年数", "O_年間基本額上限", "V_必須配置人数", "AA_評価総点", "AB_価格点割合パーセント", "AC_企画事業内容配点", "AD_職員体制実績配点", "AE_安全管理危機管理配点", "AI_プレゼン時間分", "AK_人員確保難易度"]
        for col in numeric_cols:
            if col in df.columns:
                # Remove commas, "円", etc. before converting to numeric
                df[col] = pd.to_numeric(df[col].astype(str).replace(r'[^\d.]', '', regex=True), errors='coerce')
        
        return df
    except Exception as e:
        st.error(f"データロードエラー: {e}")
        return pd.DataFrame()

def generate_ai_summary(df):
    try:
        req_texts = df["AO_提案要求リスト"].dropna().astype(str).tolist() if "AO_提案要求リスト" in df.columns else []
        prep_texts = df["AP_物理的システム的準備リスト"].dropna().astype(str).tolist() if "AP_物理的システム的準備リスト" in df.columns else []
        
        # Limit the number of texts to avoid hitting token limits
        import random
        if len(req_texts) > 100: req_texts = random.sample(req_texts, 100)
        if len(prep_texts) > 100: prep_texts = random.sample(prep_texts, 100)
        
        req_combined = "\n".join(req_texts)
        prep_combined = "\n".join(prep_texts)
        
        prompt = f"""
あなたは自治体BPO業務（学童・児童館・保育所など）の仕様書の要約システムです。
以下の「提案要求リスト」と「準備物リスト」のテキストデータを読み込み、共通要素を抽出してグループ化してください。

【出力要件】
1. 挨拶や「まとめました」などの余計な文章は絶対に含めないでください。
2. （X/226）のような数字の表記は不要です。
3. 必ず以下のJSONフォーマットのみを厳格に出力してください。マークダウンブロック(```json)も含めないでください。

[
  {{
    "title": "提案要求リストのまとめ",
    "groups": [
      {{ "group_name": "大項目名（例：運営方針・理念・事業計画）", "details": ["具体的な内容1", "具体的な内容2"] }}
    ]
  }},
  {{
    "title": "準備物リストのまとめ",
    "groups": [
      {{ "group_name": "大項目名", "details": ["具体的な内容1"] }}
    ]
  }}
]

【提案要求リストのデータ】
{req_combined}

【準備物リストのデータ】
{prep_combined}
"""
        client = genai.Client(api_key=st.secrets.get("GEMINI_API_KEY", ""))
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
        )
        res_text = response.text.strip()
        if res_text.startswith("```json"):
            res_text = res_text[7:]
        if res_text.startswith("```"):
            res_text = res_text[3:]
        if res_text.endswith("```"):
            res_text = res_text[:-3]
        return res_text.strip()
    except Exception as e:
        return f'{{"error": "AI要約の生成に失敗しました: {e}"}}'

def render_dashboard(df, spreadsheet_target):
    st.markdown("""
<style>
/* Streamlit multiselect custom height limit */
div[data-baseweb="select"] > div:first-child {
    max-height: 48px !important;
    overflow-y: auto !important;
}
</style>
""", unsafe_allow_html=True)
    if df.empty:
        st.info("📊 まだデータがありません。解析・登録からデータを追加してください。")
        return

    st.subheader("📊 施設ごとの市場データ")
    st.markdown("##### 🔍 市場データの絞り込み")
    filter_col1, filter_col2, filter_col3, filter_col4 = st.columns(4)
    
    market_df = df.copy()
    
    with filter_col1:
        if "AN_公募年度" in market_df.columns:
            years = sorted([str(x) for x in market_df["AN_公募年度"].dropna().unique() if x], reverse=True)
            selected_years = st.multiselect("公募年度", years, default=years)
            if selected_years:
                market_df = market_df[market_df["AN_公募年度"].astype(str).isin(selected_years)]
                
    with filter_col2:
        if "F_契約形態" in market_df.columns:
            forms = sorted([str(x) for x in market_df["F_契約形態"].dropna().unique() if x])
            selected_forms = st.multiselect("契約形態", forms, default=forms)
            if selected_forms:
                market_df = market_df[market_df["F_契約形態"].astype(str).isin(selected_forms)]
                
    with filter_col3:
        if "E_施設区分" in market_df.columns:
            facilities = sorted([str(x) for x in market_df["E_施設区分"].dropna().unique() if x])
            selected_facilities = st.multiselect("施設区分", facilities, default=facilities)
            if selected_facilities:
                market_df = market_df[market_df["E_施設区分"].astype(str).isin(selected_facilities)]

    with filter_col4:
        if "N_契約方式" in market_df.columns:
            methods = sorted([str(x) for x in market_df["N_契約方式"].dropna().unique() if x])
            selected_methods = st.multiselect("契約方式", methods, default=methods)
            if selected_methods:
                market_df = market_df[market_df["N_契約方式"].astype(str).isin(selected_methods)]

    st.write("---")
    
    kpi1, kpi2 = st.columns(2)
    kpi1.metric("該当する案件数", f"{len(market_df)} 件")
    
    if "O_年間基本額上限" in market_df.columns:
        avg_budget = market_df["O_年間基本額上限"].mean()
        if not pd.isna(avg_budget):
            kpi2.metric("平均予算上限", f"¥{avg_budget:,.0f}")
        else:
            kpi2.metric("平均予算上限", "-")

    # Facility Comparison
    if "E_施設区分" in market_df.columns:
        # Filter out empty facility types
        df_facilities = market_df[market_df["E_施設区分"].notna() & (market_df["E_施設区分"] != "")]
        
        if not df_facilities.empty:
            plot_df = df_facilities.copy()
            
            # Helper to create hover list
            def build_hover_list(series, max_items=5):
                lst = series.dropna().astype(str).tolist()
                if not lst: return "案件名なし"
                if len(lst) > max_items:
                    return "<br>".join([f"・{n}" for n in lst[:max_items]]) + f"<br>...ほか{len(lst)-max_items}件"
                return "<br>".join([f"・{n}" for n in lst])

            # 1. Row of Charts: Pie Chart (smaller) + Top Municipalities + Contract Length
            chart_col1, chart_col2, chart_col3 = st.columns([1, 1.2, 1.2])
            
            with chart_col1:
                st.write("**施設区分別の案件割合**")
                if "E_施設区分" in plot_df.columns:
                    grouped_pie = plot_df.groupby("E_施設区分").agg(
                        案件数=("E_施設区分", "size"),
                        案件リスト=("D_案件名", lambda x: build_hover_list(x))
                    ).reset_index()
                    fig_pie = px.pie(grouped_pie, names="E_施設区分", values="案件数", custom_data=["案件リスト"], hole=.4)
                    fig_pie.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=380, showlegend=False, hoverlabel=dict(align="left"))
                    fig_pie.update_traces(
                        textposition='inside', 
                        textinfo='percent+label',
                        hovertemplate="<b>%{label}</b><br>案件数: %{value}件<br><br>【主な案件】<br>%{customdata[0]}<extra></extra>"
                    )
                    st.plotly_chart(fig_pie, use_container_width=True)
            
            with chart_col2:
                st.write("**案件が多い地域ランキング**")
                col_m1, col_m2 = st.columns(2)
                muni_level = col_m1.radio("集計単位", ["市区町村", "都道府県"], horizontal=True, label_visibility="collapsed")
                muni_top_n = col_m2.selectbox("表示件数", [10, 20, 30], index=0, label_visibility="collapsed")
                
                if muni_level == "市区町村":
                    pref = plot_df["B_都道府県"].replace("null", "").fillna("")
                    city = plot_df["C_自治体名"].replace("null", "").fillna("")
                    plot_df["地域名"] = pref + city
                else:
                    plot_df["地域名"] = plot_df["B_都道府県"].replace("null", "").fillna("")
                    
                valid_muni = plot_df[plot_df["地域名"].str.strip() != ""]
                if not valid_muni.empty:
                    grouped_muni = valid_muni.groupby("地域名").agg(
                        案件数=("地域名", "size"),
                        案件リスト=("D_案件名", lambda x: build_hover_list(x))
                    ).reset_index()
                    top_muni = grouped_muni.sort_values("案件数", ascending=False).head(muni_top_n)
                    
                    fig_muni = px.bar(top_muni.sort_values("案件数", ascending=True), 
                                      x="案件数", y="地域名", orientation='h', custom_data=["案件リスト"])
                    fig_muni.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=380, hoverlabel=dict(align="left"))
                    fig_muni.update_traces(
                        hovertemplate="<b>%{y}</b><br>案件数: %{x}件<br><br>【主な案件】<br>%{customdata[0]}<extra></extra>"
                    )
                    st.plotly_chart(fig_muni, use_container_width=True)

            with chart_col3:
                if "M_契約期間年数" in plot_df.columns:
                    st.write("**契約年数の分布**")
                    valid_lengths = plot_df.dropna(subset=["M_契約期間年数"]).copy()
                    if not valid_lengths.empty:
                        valid_lengths["M_契約期間年数"] = valid_lengths["M_契約期間年数"].astype(int)
                        grouped_len = valid_lengths.groupby("M_契約期間年数").agg(
                            案件数=("M_契約期間年数", "size"),
                            案件リスト=("D_案件名", lambda x: build_hover_list(x))
                        ).reset_index()
                        
                        # Sort by numeric value
                        grouped_len = grouped_len.sort_values("M_契約期間年数")
                        grouped_len["契約年数ラベル"] = grouped_len["M_契約期間年数"].astype(str) + "年"
                        
                        fig_bar = px.bar(grouped_len, x="契約年数ラベル", y="案件数", custom_data=["案件リスト"])
                        fig_bar.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=380, hoverlabel=dict(align="left"))
                        fig_bar.update_traces(
                            hovertemplate="<b>%{x}</b><br>案件数: %{y}件<br><br>【主な案件】<br>%{customdata[0]}<extra></extra>"
                        )
                        st.plotly_chart(fig_bar, use_container_width=True)

            # 2. Summary Table of Averages
            st.write("**施設区分別の平均指標**")
            agg_dict = {}
            if "O_年間基本額上限" in market_df.columns:
                agg_dict["O_年間基本額上限"] = "mean"
            if "M_契約期間年数" in market_df.columns:
                agg_dict["M_契約期間年数"] = "mean"
                
            if agg_dict:
                grouped = df_facilities.groupby("E_施設区分").agg(agg_dict).reset_index()
                
                # 指定された並び順でソート（表記揺れ等も考慮して網羅的に定義）
                sort_order = ["学童", "子育て", "放課後子供教室", "複合施設", "子育て支援拠点", "児童館", "その他", "そのほか"]
                grouped["sort_key"] = grouped["E_施設区分"].apply(lambda x: sort_order.index(x) if x in sort_order else 999)
                grouped = grouped.sort_values("sort_key").drop("sort_key", axis=1)
                
                display_grouped = pd.DataFrame()
                display_grouped["施設区分"] = grouped["E_施設区分"]
                if "O_年間基本額上限" in grouped.columns:
                    display_grouped["平均 年間基本額上限(円)"] = grouped["O_年間基本額上限"].apply(lambda x: f"¥{x:,.0f}" if pd.notna(x) else "-")
                if "M_契約期間年数" in grouped.columns:
                    display_grouped["平均 契約期間年数"] = grouped["M_契約期間年数"].apply(lambda x: f"{x:.1f}年" if pd.notna(x) else "-")
                st.dataframe(display_grouped, use_container_width=True, hide_index=True)

    st.divider()
    
    st.subheader("📋 提案要求・準備物 カタログ (全体傾向)")
    st.write("過去の案件で実際に求められた事項を全データから網羅的に確認し、社内の既存アセットの適応可否を判断するための材料です。")
    
    st.write("") # Spacing
    
    cat_tab1, cat_tab2, cat_tab3 = st.tabs(["📈 要素別 出現頻度（レーダーチャート）", "🤖 要求リスト(AIまとめ)", "📑 要求事項の一覧表"])
    
    filter_df = df.copy() # Use all data, ignoring market filters
    
    with cat_tab1:
        st.write("主要なキーワードが「提案要求・準備物」の中にどれくらい出現しているかを集計します。")
        st.write("各施設区分ごとの傾向を比較できます。以下に全体と各施設区分の分布を順に表示します。")
        st.write("---")
        
        if "AO_提案要求リスト" in df.columns or "AP_物理的システム的準備リスト" in df.columns:
            keywords = {
                "安全・防災（危機管理）": ["安全", "防災", "危機管理", "避難", "事故"],
                "運営・マネジメント": ["運営", "管理", "方針", "理念", "体制"],
                "学習・プログラム": ["学習", "プログラム", "指導", "教育", "イベント"],
                "食育・おやつ": ["食", "おやつ", "アレルギー", "提供"],
                "ICT・システム": ["ICT", "システム", "アプリ", "タブレット", "導入"],
                "地域連携・保護者": ["地域", "保護者", "連携", "コミュニケーション"],
                "研修・人材育成": ["研修", "育成", "人材", "スキル"]
            }
            
            if "E_施設区分" in filter_df.columns:
                radar_facilities = sorted([str(x) for x in filter_df["E_施設区分"].dropna().unique() if x])
                sort_order = ["学童", "子育て", "放課後子供教室", "複合施設", "子育て支援拠点", "児童館", "その他", "そのほか"]
                radar_facilities = sorted(radar_facilities, key=lambda x: sort_order.index(x) if x in sort_order else 999)
            else:
                radar_facilities = []
                
            loop_targets = [("全体（合計）", filter_df.copy())]
            for fac in radar_facilities:
                loop_targets.append((fac, filter_df[filter_df["E_施設区分"].astype(str) == fac].copy()))
                
            has_data = False
            for fac_name, target_df in loop_targets:
                total_records = len(target_df)
                if total_records == 0:
                    continue
                    
                has_data = True
                fit_stats = []
                text_series = pd.Series(dtype=str)
                if "AO_提案要求リスト" in target_df.columns:
                    text_series = target_df["AO_提案要求リスト"].astype(str)
                if "AP_物理的システム的準備リスト" in target_df.columns:
                    text_series = text_series + " " + target_df["AP_物理的システム的準備リスト"].astype(str)
                    
                for label, kws in keywords.items():
                    count = text_series.str.contains("|".join(kws), case=False, na=False).sum()
                    fit_stats.append({"要素": label, "出現件数": count, "全体に対する割合(%)": round((count/total_records)*100, 1)})
                
                fit_df = pd.DataFrame(fit_stats)
                
                st.markdown(f"#### 【{fac_name}】の要求キーワード分布")
                col_chart1, col_chart2 = st.columns(2)
                with col_chart1:
                    plot_df = fit_df.copy()
                    plot_df.columns = ["要素", "出現件数", "割合"]
                    
                    fig_radar = px.line_polar(plot_df, r="割合", theta="要素", line_close=True, markers=True,
                                              hover_data=["出現件数"],
                                              range_r=[0, 100])
                    fig_radar.update_traces(fill='toself')
                    st.plotly_chart(fig_radar, use_container_width=True)
                with col_chart2:
                    st.dataframe(fit_df, use_container_width=True, hide_index=True)
                
                st.write("---")
                
            if not has_data:
                st.warning("対象となるデータがありません。")

    with cat_tab2:
        CACHE_PATH = "ai_catalog_summary_cache_v2.json"
        
        # Adjust column ratio to bring the right side further left
        col_left, col_right = st.columns([1, 4])
        with col_left:
            st.write("AIを利用して全案件の要求をまとめます。")
            if st.button("🚀 まとめを生成・更新する"):
                with st.spinner("AIが要約を生成しています... (数十秒かかります)"):
                    summary_result = generate_ai_summary(filter_df)
                    try:
                        with open(CACHE_PATH, "w", encoding="utf-8") as f:
                            f.write(summary_result)
                        st.success("まとめを更新しました！")
                    except Exception as err:
                        st.error(f"保存に失敗しました: {err}")
        
        with col_right:
            if os.path.exists(CACHE_PATH):
                with open(CACHE_PATH, "r", encoding="utf-8") as f:
                    cached_text = f.read()
                try:
                    import json
                    data = json.loads(cached_text)
                    if isinstance(data, dict) and "error" in data:
                        st.error(data["error"])
                    else:
                        for section in data:
                            st.markdown(f"#### {section.get('title', '')}")
                            for group in section.get('groups', []):
                                with st.expander(group.get('group_name', '無題')):
                                    for detail in group.get('details', []):
                                        st.write(f"・{detail}")
                except Exception as e:
                    st.info(cached_text)
            else:
                st.info("💡 左側のボタンを押してまとめを生成してください。")

    with cat_tab3:
        st.write(f"**詳細データ一覧 ({len(filter_df)}件)**")
        display_cols = ["C_自治体名", "D_案件名", "E_施設区分", "AO_提案要求リスト", "AP_物理的システム的準備リスト"]
        available_cols = [c for c in display_cols if c in df.columns]
        
        display_raw_df = filter_df[available_cols].copy()
        display_raw_df = display_raw_df.rename(columns=DISPLAY_COLS)
        st.dataframe(display_raw_df, use_container_width=True, hide_index=True)

# --- Streamlit UI ---
def main():
    st.set_page_config(page_title="自治体BPO案件解析システム", layout="wide", page_icon="logo.png")
    
    st.markdown("""
        <style>
            h1 { color: #333333; font-weight: 300; }
            .stButton>button { background-color: #4A90E2; color: white; border-radius: 4px; border: none; padding: 0.5em 2em; }
        </style>
    """, unsafe_allow_html=True)

    col_title1, col_title2 = st.columns([1, 10])
    with col_title1:
        if os.path.exists("logo.png"):
            st.image("logo.png", width=80)
    with col_title2:
        st.title("自治体BPO案件解析システム")
    st.write("資料をアップロード → AIが解析 → スプレッドシートに自動登録")

    # --- Session State Init ---
    if "phase" not in st.session_state:
        st.session_state.phase = "upload"  # upload → scanning → group_confirm → extracting → confirm → writing
    if "groups" not in st.session_state:
        st.session_state.groups = []
    if "temp_paths" not in st.session_state:
        st.session_state.temp_paths = []
    if "extracted_results" not in st.session_state:
        st.session_state.extracted_results = []
    if "summaries" not in st.session_state:
        st.session_state.summaries = []
    if "original_groups" not in st.session_state:
        st.session_state.original_groups = []
    if "chk_generation" not in st.session_state:
        st.session_state.chk_generation = 0

    # Secrets loading
    api_key = None
    spreadsheet_target = "自治体案件管理"
    try:
        if len(st.secrets) > 0:
            api_key = st.secrets.get("GEMINI_API_KEY")
            spreadsheet_target = st.secrets.get("GSHEET_ID") or st.secrets.get("GSHEET_NAME", "自治体案件管理")
    except Exception:
        pass

    if not api_key:
        with st.sidebar:
            st.header("設定")
            api_key = st.text_input("Gemini API Key", type="password")
            spreadsheet_target = st.text_input("Googleスプレッドシート名 または ID", value=spreadsheet_target)
            st.divider()
            st.info("Streamlit Cloudでデプロイする場合は Secrets に登録してください。")

    tab_main, tab_dash = st.tabs(["📄 解析・登録", "📊 分析ダッシュボード"])

    with tab_main:
        # ========================================
        # PHASE 1: Upload & Mode Selection
        # ========================================
        uploaded_files = st.file_uploader(
            "📎 資料をすべてドロップ（PDF, Word, Excel, 画像 等）",
            accept_multiple_files=True
        )

        mode = st.radio(
            "アップロードモード",
            ["📦 モードA: すべて同じ案件です", "🔀 モードB: 複数の案件が混ざっています（AIが自動分類）"],
            index=0,
            help="モードA: 全ファイルを1つの案件として処理します。モードB: AIが案件ごとに自動分類します。"
        )
        is_mode_a = "モードA" in mode

        # Start button (only shown in upload phase)
        if st.session_state.phase == "upload":
            # Form-like wrapper to prevent accidental submissions
            with st.form("upload_form"):
                submit_button = st.form_submit_button("🚀 解析を開始")
                
            if submit_button and uploaded_files and api_key:
                st.session_state.phase = "scanning"
                st.session_state.groups = []
                st.session_state.original_groups = []
                st.session_state.extracted_results = []
                st.session_state.summaries = []
                
                # Clean up old temp files
                for p in st.session_state.temp_paths:
                    if os.path.exists(p):
                        os.remove(p)
                st.session_state.temp_paths = []
                
                # Deduplicate & save files immediately upon button click
                temp_paths = []
                seen_hashes = set()
                skipped = []
                
                for i, uf in enumerate(uploaded_files):
                    fh = hashlib.md5(uf.getvalue()).hexdigest()
                    if fh in seen_hashes:
                        skipped.append(uf.name)
                        continue
                    seen_hashes.add(fh)
                    ext = os.path.splitext(uf.name)[1]
                    tp = f"temp_file_{i}{ext}"
                    with open(tp, "wb") as f:
                        f.write(uf.getbuffer())
                    temp_paths.append(tp)
                
                st.session_state.temp_paths = temp_paths
                if skipped:
                    st.session_state.skipped_files = skipped
                else:
                    st.session_state.skipped_files = []
                    
                st.rerun()

        # ========================================
        # PHASE 1.1: Scanning & Grouping
        # ========================================
        if st.session_state.phase == "scanning":
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            temp_paths = st.session_state.temp_paths
            
            if getattr(st.session_state, "skipped_files", []):
                st.info(f"ℹ️ 重複スキップ: {', '.join(st.session_state.skipped_files)}")
                
            if not temp_paths:
                st.warning("有効なファイルがありませんでした。")
                if st.button("⬅️ アップロード画面に戻る"):
                    st.session_state.phase = "upload"
                    st.rerun()
                st.stop()
            
            progress_bar.progress(0.05)
            
            if is_mode_a:
                # MODE A: Skip scanning & grouping, treat all as one project
                status_text.text("📦 モードA: すべて同じ案件として準備中...")
                st.session_state.groups = [{
                    "group_name": f"アップロード案件（{len(temp_paths)}ファイル）",
                    "municipality": "",
                    "project_name": "",
                    "fiscal_year": "",
                    "file_indices": list(range(len(temp_paths)))
                }]
                progress_bar.progress(1.0)
                status_text.text("準備完了しました。抽出に進みます。")
                st.session_state.phase = "extracting"
                st.rerun()
            else:
                # MODE B: Scan & Group
                summaries = step1_scan_files(api_key, temp_paths, 
                                             status_callback=lambda msg: status_text.text(msg))
                st.session_state.summaries = summaries
                progress_bar.progress(0.2)
                
                with st.expander("📋 スキャン結果", expanded=False):
                    for s in summaries:
                        st.write(f"**ファイル{s['file_index']}**: {s.get('municipality', '?')} / {s.get('project_name', '?')} / {s.get('fiscal_year', '?')} — {s.get('summary', '')}")
                
                if len(temp_paths) == 1:
                    st.session_state.groups = [{
                        "group_name": f"{summaries[0].get('municipality', '')} {summaries[0].get('project_name', '')} ({summaries[0].get('fiscal_year', '')})",
                        "municipality": summaries[0].get("municipality", ""),
                        "project_name": summaries[0].get("project_name", ""),
                        "fiscal_year": summaries[0].get("fiscal_year", ""),
                        "file_indices": [0]
                    }]
                    progress_bar.progress(1.0)
                    st.session_state.phase = "extracting"
                    st.rerun()
                else:
                    status_text.text("🗂️ AIが案件ごとにグループ分け中...")
                    groups_result = step2_group_files(api_key, summaries)
                    st.session_state.groups = groups_result.get("groups", [])
                    st.session_state.original_groups = copy.deepcopy(st.session_state.groups)
                    progress_bar.progress(1.0)
                    status_text.text("グループ分け完了しました。確認してください。")
                    st.session_state.phase = "group_confirm"
                    st.rerun()

        # ========================================
        # PHASE 1.5: Group Confirm (Manual Merge)
        # ========================================
        if st.session_state.phase == "group_confirm":
            st.divider()
            st.subheader("🗂️ AIのグループ分け結果の確認")
            st.write("同じ案件の資料はチェックを入れて「選択した案件を統合する」を押してください。")
            
            # Display current groups with checkboxes
            selected_groups = []
            for gi, g in enumerate(st.session_state.groups):
                filenames = []
                for idx in g["file_indices"]:
                    # Match index to original summary
                    for s in st.session_state.summaries:
                        if s["file_index"] == idx:
                            filenames.append(f"- ファイル{idx}: {s.get('summary', '')}")
                            break
                
                with st.container(border=True):
                    col1, col2 = st.columns([1, 10])
                    with col1:
                        is_selected = st.checkbox("選択", label_visibility="collapsed", key=f"grp_chk_{gi}_{st.session_state.chk_generation}")
                        if is_selected:
                            selected_groups.append(gi)
                    with col2:
                        st.write(f"**{g['group_name']}** ({len(g['file_indices'])}ファイル)")
                        for fn in filenames:
                            st.caption(fn)

            # Merge Actions
            col_action1, col_action2 = st.columns([1, 3])
            with col_action1:
                if st.button("🔗 選択した案件を統合する"):
                    if len(selected_groups) >= 2:
                        # Merge logic
                        new_groups = []
                        merged_indices = []
                        merged_name = st.session_state.groups[selected_groups[0]]["group_name"] + " (統合)"
                        
                        for i, g in enumerate(st.session_state.groups):
                            if i in selected_groups:
                                merged_indices.extend(g["file_indices"])
                            else:
                                new_groups.append(g)
                        
                        # Add newly merged group
                        new_groups.append({
                            "group_name": merged_name,
                            "municipality": st.session_state.groups[selected_groups[0]].get("municipality", ""),
                            "project_name": st.session_state.groups[selected_groups[0]].get("project_name", ""),
                            "fiscal_year": st.session_state.groups[selected_groups[0]].get("fiscal_year", ""),
                            "file_indices": sorted(list(set(merged_indices)))
                        })
                        
                        st.session_state.groups = new_groups
                        
                        # Force checkboxes to uncheck by changing their key generation
                        st.session_state.chk_generation += 1
                                
                        st.rerun()
                    else:
                        st.warning("統合するには2つ以上の案件を選択してください。")
            with col_action2:
                if st.button("🔄 グループ分けを最初（AIの提案）に戻す"):
                    st.session_state.groups = copy.deepcopy(st.session_state.original_groups)
                    # Force checkboxes to uncheck
                    st.session_state.chk_generation += 1
                    st.rerun()
            
            st.write("---")
            if st.button("🚀 このグループで本格抽出を開始する", type="primary"):
                st.session_state.phase = "extracting"
                st.rerun()
                
        # ========================================
        # PHASE 1.8: Extracting Data
        # ========================================
        if st.session_state.phase == "extracting":
            st.divider()
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Check existing data
            with st.status("既存データとの照合中...", expanded=False) as status:
                sheet = get_sheet(st.secrets["GSHEET_ID"])
                annotated = step3_check_existing(api_key, sheet, st.session_state.groups)
                status.update(label="照合完了！", state="complete")
            
            progress_bar.progress(0.1)
            
            # Extract data for each group
            all_results = []
            for gi, group in enumerate(annotated):
                gname = group.get("group_name", f"案件{gi+1}")
                file_indices = group.get("file_indices", [])
                gpaths = [st.session_state.temp_paths[idx] for idx in file_indices if idx < len(st.session_state.temp_paths)]
                
                if not gpaths:
                    continue
                
                status_text.text(f"🧠 解析中: [{gi+1}/{len(annotated)}] {gname}...")
                
                existing_data = None
                if group["mode"] == "update" and group.get("existing_row"):
                    try:
                        row_vals = sheet.row_values(group["existing_row"])
                        existing_data = {}
                        for ci, col in enumerate(COLUMN_LAYOUT):
                            if ci < len(row_vals) and row_vals[ci]:
                                existing_data[col] = row_vals[ci]
                    except Exception:
                        pass
                
                try:
                    extracted = step4_extract_data(api_key, gpaths, existing_data)
                    all_results.append({
                        "group": group,
                        "data": extracted,
                        "group_name": gname
                    })
                except Exception as e:
                    st.error(f"❌ {gname} の解析エラー: {str(e)}")
                
                progress_bar.progress(0.1 + (0.9 * (gi + 1) / len(annotated)))
            
            st.session_state.extracted_results = all_results
            st.session_state.phase = "confirm"
            status_text.text("解析完了！以下のプレビューを確認してください。")
            st.rerun()

        # ========================================
        # PHASE 2: Preview & Confirm
        # ========================================
        if st.session_state.phase == "confirm" and st.session_state.extracted_results:
            st.divider()
            st.subheader("📊 抽出データのプレビュー")
            st.write("登録する案件にチェックを入れて「登録する」ボタンを押してください。")
            
            selected_results_to_write = []
            for ri, result in enumerate(st.session_state.extracted_results):
                group = result["group"]
                data = result["data"]
                gname = result["group_name"]
                mode_label = "🔄 更新" if group["mode"] == "update" else "🆕 新規"
                row_label = f"（行 {group['existing_row']}）" if group.get("existing_row") else ""
                
                with st.container(border=True):
                    c_chk, c_exp = st.columns([1, 15])
                    with c_chk:
                        do_register = st.checkbox("登録", value=True, key=f"reg_chk_{ri}", label_visibility="collapsed")
                        if do_register:
                            selected_results_to_write.append(result)
                    with c_exp:
                        with st.expander(f"{mode_label} {gname} {row_label}", expanded=False):
                            # Show data as a clean table
                            preview_data = {}
                            for col in COLUMN_LAYOUT:
                                if col in ("A_No", "G_ステータス"):
                                    continue
                                display_name = DISPLAY_COLS.get(col, col)
                                val = data.get(col, "")
                                if val is not None and val != "":
                                    preview_data[display_name] = str(val)
                            
                            if preview_data:
                                col1, col2 = st.columns(2)
                                items = list(preview_data.items())
                                mid = (len(items) + 1) // 2
                                with col1:
                                    for k, v in items[:mid]:
                                        st.write(f"**{k}**: {v}")
                                with col2:
                                    for k, v in items[mid:]:
                                        st.write(f"**{k}**: {v}")
                            else:
                                st.warning("データが抽出されませんでした。")
            
            col_confirm, col_cancel = st.columns(2)
            with col_confirm:
                if st.button("✅ チェックした内容で登録する", type="primary"):
                    if not selected_results_to_write:
                        st.warning("登録する案件を少なくとも1つ選択してください。")
                    else:
                        st.session_state.final_selection = selected_results_to_write
                        st.session_state.phase = "writing"
                        st.rerun()
            with col_cancel:
                if st.button("❌ キャンセル"):
                    # Cleanup
                    for p in st.session_state.temp_paths:
                        if os.path.exists(p):
                            os.remove(p)
                    st.session_state.phase = "upload"
                    st.session_state.extracted_results = []
                    st.session_state.temp_paths = []
                    st.rerun()

        # ========================================
        # PHASE 3: Write to Sheet
        # ========================================
        if st.session_state.phase == "writing":
            st.divider()
            status_text = st.empty()
            status_text.text("📝 スプレッドシートへ書き込み中...")
            
            try:
                sheet = get_sheet(spreadsheet_target)
                all_values = sheet.col_values(1)
                next_no = len([v for v in all_values[1:] if str(v).strip()]) + 1
                
                success_count = 0
                to_process = getattr(st.session_state, "final_selection", st.session_state.extracted_results)
                
                for result in to_process:
                    group = result["group"]
                    data = result["data"]
                    gname = result["group_name"]
                    
                    row = build_row(data, next_no, group["mode"], sheet, group.get("existing_row"))
                    
                    try:
                        step5_write_to_sheet(sheet, row, group["mode"], group.get("existing_row"))
                        if group["mode"] == "update":
                            st.success(f"🔄 {gname} → 行 {group['existing_row']} を更新しました")
                        else:
                            st.success(f"🆕 {gname} → 行 {next_no} に新規登録しました")
                            next_no += 1
                        success_count += 1
                    except Exception as e:
                        st.error(f"❌ {gname} の書き込みエラー: {str(e)}")
                
                if success_count > 0:
                    st.balloons()
                    st.success(f"🎉 合計 {success_count} 件の処理が完了しました！")
                    st.cache_data.clear()
            
            except Exception as e:
                st.error(f"❌ エラー: {str(e)}")
            finally:
                for p in st.session_state.temp_paths:
                    if os.path.exists(p):
                        os.remove(p)
                st.session_state.phase = "upload"
                st.session_state.extracted_results = []
                st.session_state.temp_paths = []

    with tab_dash:
        if api_key:
            with st.spinner("データを読み込み中..."):
                df = load_historical_data(spreadsheet_target)
            render_dashboard(df, spreadsheet_target)
        else:
            st.warning("APIキーとスプレッドシートの設定を完了してください。")

if __name__ == "__main__":
    main()
